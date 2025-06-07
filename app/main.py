# app/main.py
"""
Aplikacja FastAPI do zarządzania opiniami sądowymi.
Refaktoryzowany główny plik - zawiera tylko konfigurację i kluczowe endpointy.
"""

import asyncio
import time
import uuid
import shutil
from pathlib import Path

from fastapi import FastAPI, Request, UploadFile, File, Form, HTTPException
from fastapi.responses import RedirectResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from sqlmodel import Session, select
from datetime import datetime

# Lokalne importy
from app.db import engine, FILES_DIR, BASE_DIR, init_db
from app.models import Document
from app.document_utils import (
    detect_mime_type, 
    check_file_extension, 
    get_content_type_from_mime, 
    STEP_ICON
)
from app.text_extraction import get_text_preview
from app.routes import opinions_router, documents_router, updates_router

# Import dla OCR
from tasks.ocr_manager import enqueue

# Konfiguracja aplikacji
app = FastAPI(title="Court Workflow", description="System zarządzania opiniami sądowymi")

# Konfiguracja statycznych plików i szablonów
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))
app.mount("/files", StaticFiles(directory=str(FILES_DIR)), name="files")
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")

# Rejestracja routerów
app.include_router(opinions_router)
app.include_router(documents_router)
app.include_router(updates_router)

@app.on_event("startup")
async def startup():
    """Inicjalizacja aplikacji przy starcie."""
    # Ensure database tables are created
    init_db()
    
    # Uruchomienie nowego systemu workerów zadań w tle
    from app.background_tasks import start_background_workers
    asyncio.create_task(start_background_workers())

# Middleware do monitorowania wydajności
@app.middleware("http")
async def detect_blocking_operations(request: Request, call_next):
    """Middleware do wykrywania blokujących operacji."""
    # Rozpocznij liczenie czasu
    start_time = time.time()
    
    # Zidentyfikuj żądanie
    request_id = str(uuid.uuid4())[:8]
    path = request.url.path
    print(f"[{request_id}] Rozpoczęto żądanie: {path}")
    
    # Wykonaj żądanie
    response = await call_next(request)
    
    # Sprawdź czas wykonania
    elapsed = time.time() - start_time
    print(f"[{request_id}] Zakończono żądanie: {path} w {elapsed:.4f}s")
    
    # Loguj długie żądania
    if elapsed > 1.0:
        print(f"[{request_id}] UWAGA: Długie żądanie ({elapsed:.4f}s): {path}")
    
    return response

# ==================== ENDPOINTY UPLOAD ====================

@app.get("/upload")
def upload_form(request: Request):
    """Formularz uploadowania nowych opinii."""
    from app.document_utils import ALLOWED_EXTENSIONS
    allowed_types = ", ".join(ALLOWED_EXTENSIONS.keys())
    return templates.TemplateResponse(
        "upload.html", 
        {"request": request, "title": "Załaduj nową opinię", "allowed_types": allowed_types}
    )

@app.post("/upload")
async def upload(request: Request, files: list[UploadFile] = File(...)):
    """Dodawanie nowych głównych dokumentów (opinii)."""
    uploaded_docs = []
    
    with Session(engine) as session:
        for file in files:
            # Sprawdzenie rozszerzenia pliku
            suffix = check_file_extension(file.filename)
            
            # Dla opinii akceptujemy tylko pliki Word
            if suffix.lower() not in ['.doc', '.docx']:
                raise HTTPException(
                    status_code=400,
                    detail=f"Opinie muszą być w formacie Word (.doc, .docx). Przesłano: {suffix}"
                )
            
            # Generowanie unikalnej nazwy pliku
            unique_name = f"{uuid.uuid4().hex}{suffix}"
            dest = FILES_DIR / unique_name
            
            # Zapisanie pliku
            with dest.open("wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            
            # Wykrywanie właściwego MIME typu pliku
            actual_mime_type = detect_mime_type(dest)
            
            # Zapisanie do bazy danych jako dokument główny
            doc = Document(
                original_filename=file.filename,
                stored_filename=unique_name,
                step="k1",  # Nowe opinie zaczynają od k1
                ocr_status="none",  # Word nie wymaga OCR
                is_main=True,  # Oznacz jako dokument główny
                content_type="opinion",
                mime_type=actual_mime_type,
                doc_type="Opinia",
                creator=None  # Tu można dodać current_user
            )
            session.add(doc)
            session.commit()
            uploaded_docs.append(doc.id)

    # Przekieruj do listy opinii lub do szczegółów pierwszej dodanej opinii
    if len(uploaded_docs) == 1:
        return RedirectResponse(request.url_for("opinion_detail", doc_id=uploaded_docs[0]), status_code=303)
    else:
        return RedirectResponse(request.url_for("list_opinions"), status_code=303)

# ==================== ENDPOINTY TWORZENIA OPINII ====================

@app.get("/create_empty_opinion")
def create_empty_opinion_form(request: Request):
    """Formularz tworzenia nowej pustej opinii."""
    return templates.TemplateResponse(
        "create_empty_opinion.html", 
        {"request": request, "title": "Utwórz nową pustą opinię"}
    )

@app.post("/create_empty_opinion")
def create_empty_opinion(
    request: Request,
    sygnatura: str | None = Form(None),
    doc_type: str = Form(...),
    step: str = Form("k1")
):
    """Utworzenie nowej pustej opinii bez dokumentu."""
    # Generowanie unikalnej nazwy dla "pustego" dokumentu
    unique_name = f"{uuid.uuid4().hex}.empty"
    
    with Session(engine) as session:
        # Utworzenie nowej opinii w bazie danych
        opinion = Document(
            original_filename="Nowa opinia",
            stored_filename=unique_name,
            step=step,
            ocr_status="none",
            is_main=True,
            content_type="opinion",
            doc_type=doc_type,
            sygnatura=sygnatura,
            creator=None  # Tu można dodać current_user gdy będzie system użytkowników
        )
        session.add(opinion)
        session.commit()
        opinion_id = opinion.id
    
    # Przekieruj do widoku szczegółów opinii
    return RedirectResponse(request.url_for("opinion_detail", doc_id=opinion_id), status_code=303)

# ==================== ENDPOINTY SZYBKIEGO OCR ====================

@app.get("/quick_ocr")
def quick_ocr_form(request: Request):
    """Formularz do szybkiego OCR dokumentów bez przypisywania do opinii."""
    from app.document_utils import ALLOWED_EXTENSIONS
    allowed_types = ", ".join([k for k in ALLOWED_EXTENSIONS.keys() 
                              if k not in ['.doc', '.docx']])  # Bez plików Word
    return templates.TemplateResponse(
        "quick_ocr.html", 
        {"request": request, "title": "Szybki OCR", "allowed_types": allowed_types}
    )

@app.post("/quick_ocr")
async def quick_ocr(request: Request, files: list[UploadFile] = File(...)):
    """Szybki OCR - dodawanie dokumentów bez wiązania z opinią."""
    uploaded_docs = []
    
    # Utwórz lub pobierz specjalną "opinię" dla dokumentów niezwiązanych
    with Session(engine) as session:
        # Sprawdź czy istnieje specjalna opinia dla dokumentów niezwiązanych
        special_opinion_query = select(Document).where(
            Document.is_main == True,
            Document.doc_type == "Dokumenty niezwiązane z opiniami"
        )
        special_opinion = session.exec(special_opinion_query).first()
        
        # Jeśli nie istnieje, utwórz ją
        if not special_opinion:
            special_opinion = Document(
                original_filename="Dokumenty niezwiązane z opiniami",
                stored_filename=f"{uuid.uuid4().hex}.empty",
                step="k1",
                ocr_status="none",
                is_main=True,
                content_type="container",  # Specjalny typ dla kontenera dokumentów
                doc_type="Dokumenty niezwiązane z opiniami",
                sygnatura="UNASSIGNED",
                creator=None
            )
            session.add(special_opinion)
            session.commit()
            special_opinion_id = special_opinion.id
        else:
            special_opinion_id = special_opinion.id
    
    # Przetwarzanie wgranych plików
    for file in files:
        # Sprawdzenie rozszerzenia pliku
        suffix = check_file_extension(file.filename)
        
        # Ignorujemy pliki Word w szybkim OCR
        if suffix.lower() in ['.doc', '.docx']:
            continue
        
        # Generowanie unikalnej nazwy pliku
        unique_name = f"{uuid.uuid4().hex}{suffix}"
        dest = FILES_DIR / unique_name
        
        # Zapisanie pliku
        content = await file.read()
        with open(dest, "wb") as buffer:
            buffer.write(content)
        
        # Wykrywanie właściwego MIME typu pliku
        actual_mime_type = detect_mime_type(dest)
        
        # Określanie content_type na podstawie MIME type
        content_type = get_content_type_from_mime(actual_mime_type)
        
        # Ustal właściwy status OCR
        ocr_status = "pending"
        
        # Zapisanie do bazy danych
        with Session(engine) as session:
            new_doc = Document(
                doc_type="Dokument OCR",
                original_filename=file.filename,
                stored_filename=unique_name,
                step="k1",
                ocr_status=ocr_status,
                parent_id=special_opinion_id,  # Przypisz do specjalnej "opinii"
                is_main=False,
                content_type=content_type,
                mime_type=actual_mime_type,
                creator=None,
                upload_time=datetime.now()
            )
            session.add(new_doc)
            session.commit()
            uploaded_docs.append(new_doc.id)
    
    # Uruchom OCR dla wgranych dokumentów
    asyncio.create_task(_enqueue_ocr_documents_nonblocking(uploaded_docs))
    
    # Przekieruj do listy dokumentów
    return RedirectResponse(request.url_for("list_documents"), status_code=303)

# ==================== ENDPOINTY UPLOADU DO OPINII ====================

@app.get("/opinion/{doc_id}/upload")
def opinion_upload_form(request: Request, doc_id: int):
    """Formularz dodawania dokumentów do opinii."""
    with Session(engine) as session:
        opinion = session.get(Document, doc_id)
        if not opinion or not opinion.is_main:
            raise HTTPException(status_code=404, detail="Nie znaleziono opinii")
    
    from app.document_utils import ALLOWED_EXTENSIONS
    allowed_types = ", ".join(ALLOWED_EXTENSIONS.keys())
    return templates.TemplateResponse(
        "upload_to_opinion.html", 
        {
            "request": request, 
            "opinion": opinion,
            "allowed_types": allowed_types,
            "title": f"Dodaj dokumenty do opinii: {opinion.sygnatura or opinion.original_filename}"
        }
    )

@app.post("/opinion/{doc_id}/upload")
async def opinion_upload(request: Request, doc_id: int, 
                         doc_type: str = Form(...),
                         files: list[UploadFile] = File(...),
                         run_ocr: bool = Form(False)):
    """Dodawanie dokumentów do opinii."""
    # Sprawdź czy opinia istnieje
    with Session(engine) as session:
        opinion = session.get(Document, doc_id)
        if not opinion or not opinion.is_main:
            raise HTTPException(status_code=404, detail="Nie znaleziono opinii")
    
    uploaded_docs = []
    has_ocr_docs = False  # Flaga wskazująca, czy jakiekolwiek dokumenty wymagają OCR

    # Przetwarzanie wgranych plików
    for file in files:
        # Sprawdzenie rozszerzenia pliku
        suffix = check_file_extension(file.filename)
        
        # Generowanie unikalnej nazwy pliku
        unique_name = f"{uuid.uuid4().hex}{suffix}"
        dest = FILES_DIR / unique_name
        
        # Zapisanie pliku
        content = await file.read()
        
        with open(dest, "wb") as buffer:
            buffer.write(content)
        
        # Wykrywanie właściwego MIME typu pliku
        actual_mime_type = detect_mime_type(dest)
        
        # Określanie content_type na podstawie MIME type
        content_type = get_content_type_from_mime(actual_mime_type)
        
        # Jeśli to nowy dokument główny, nie powiązuj go z obecną opinią
        is_main = content_type == "opinion" and doc_type == "Opinia"
        parent_id = None if is_main else doc_id
        
        # Ustal właściwy status OCR
        ocr_status = "none"
        if run_ocr and content_type != "opinion":
            ocr_status = "pending"
            has_ocr_docs = True  # Zaznacz, że co najmniej jeden dokument wymaga OCR
        
        # Zapisanie do bazy danych
        with Session(engine) as session:
            new_doc = Document(
                sygnatura=opinion.sygnatura,
                doc_type=doc_type,
                original_filename=file.filename,
                stored_filename=unique_name,
                step="k1" if is_main else opinion.step,  # Nowe opinie zaczynają od k1
                ocr_status=ocr_status,
                parent_id=parent_id,
                is_main=is_main,
                content_type=content_type,
                mime_type=actual_mime_type,
                creator=None,  # Tu można dodać current_user
                upload_time=datetime.now()
            )
            session.add(new_doc)
            session.commit()
            
            uploaded_docs.append(new_doc.id)

    # Uruchom OCR dla wgranych dokumentów w tle
    if has_ocr_docs:
        asyncio.create_task(_enqueue_ocr_documents_nonblocking(uploaded_docs))

    # Przekieruj do widoku opinii z odpowiednim komunikatem
    redirect_url = request.url_for("opinion_detail", doc_id=doc_id)

    if has_ocr_docs:
        return RedirectResponse(
            f"{redirect_url}?ocr_started=true&count={len(uploaded_docs)}", 
            status_code=303
        )
    else:
        return RedirectResponse(redirect_url, status_code=303)

# ==================== ENDPOINTY OCR ====================

@app.post("/document/{doc_id}/run_ocr")
async def document_run_ocr(request: Request, doc_id: int):
    """Endpoint do ręcznego uruchomienia OCR ponownie."""
    with Session(engine) as session:
        doc = session.get(Document, doc_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Nie ma takiego dokumentu")
        doc.ocr_status = "pending"
        doc.ocr_progress = 0.0
        doc.ocr_progress_info = "Oczekuje w kolejce"
        session.add(doc)
        session.commit()
    
    # Dodaj do kolejki OCR
    from app.background_tasks import enqueue_ocr_task
    asyncio.create_task(enqueue_ocr_task(doc_id))
    
    # Dodaj parametr do URL przekierowania, aby pokazać powiadomienie
    redirect_url = request.url_for("document_detail", doc_id=doc_id)
    return RedirectResponse(f"{redirect_url}?ocr_restarted=true", status_code=303)

@app.get("/api/document/{doc_id}/ocr-progress")
def document_ocr_progress(doc_id: int):
    """Zwraca informacje o postępie OCR w formacie JSON."""
    with Session(engine) as session:
        doc = session.get(Document, doc_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Nie znaleziono dokumentu")
        
        # Przygotuj dane o postępie
        progress_data = {
            "status": doc.ocr_status,
            "progress": doc.ocr_progress or 0.0,
            "info": doc.ocr_progress_info or "",
            "current_page": doc.ocr_current_page or 0,
            "total_pages": doc.ocr_total_pages or 0,
            "confidence": doc.ocr_confidence
        }
        
        return progress_data

# ==================== ENDPOINTY PODGLĄDU ====================

@app.get("/document/{doc_id}/preview-content")
def document_preview_content(request: Request, doc_id: int):
    """Zwraca HTML z zawartością podglądu dokumentu do wyświetlenia w modalu."""
    with Session(engine) as session:
        doc = session.get(Document, doc_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Nie ma takiego dokumentu")
        
        # Sprawdź, czy istnieje wynik OCR (dokument TXT)
        ocr_txt = None
        if doc.ocr_status == "done":
            ocr_txt_query = select(Document).where(
                Document.ocr_parent_id == doc_id,
                Document.doc_type == "OCR TXT"
            )
            ocr_txt = session.exec(ocr_txt_query).first()
    
    # Tu powinna być długa funkcja generująca HTML podglądu
    # Można ją wynieść do osobnego modułu jeśli będzie za długa
    return _generate_preview_html(request, doc, ocr_txt)

# ==================== ENDPOINTY PDF VIEWER ====================

@app.get("/document/{doc_id}/pdf-viewer")
def document_pdf_viewer(request: Request, doc_id: int):
    """Zaawansowany podgląd PDF z funkcją zaznaczania i OCR."""
    with Session(engine) as session:
        doc = session.get(Document, doc_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Nie znaleziono dokumentu")
        
        # Sprawdź czy dokument to PDF
        if not doc.mime_type or doc.mime_type != 'application/pdf':
            raise HTTPException(status_code=400, detail="Ten widok jest dostępny tylko dla plików PDF")
    
    return templates.TemplateResponse(
        "pdf_view_with_selection.html", 
        {
            "request": request, 
            "doc": doc,
            "title": f"Podgląd PDF z zaznaczaniem - {doc.original_filename}"
        }
    )

@app.post("/api/document/{doc_id}/ocr-selection")
async def document_ocr_selection(request: Request, doc_id: int):
    """Zwraca OCR dla zaznaczonego fragmentu dokumentu (PDF lub obraz)."""
    # Import funkcji OCR
    from tasks.ocr.models import process_image_to_text
    import PyPDF2
    from pdf2image import convert_from_path
    import tempfile
    from PIL import Image
    from tasks.ocr.config import logger
    
    try:
        # Pobierz dane z POST
        data = await request.json()
        page = data.get('page', 1)     # Numer strony (1-based dla PDF, zawsze 1 dla obrazów)
        x1 = data.get('x1', 0)         # Współrzędne zaznaczenia (0-1)
        y1 = data.get('y1', 0)
        x2 = data.get('x2', 1)
        y2 = data.get('y2', 1)
        skip_pdf_embed = data.get('skip_pdf_embed', False)
        
        # Sprawdź czy dokument istnieje
        with Session(engine) as session:
            doc = session.get(Document, doc_id)
            if not doc:
                return {"error": "Nie znaleziono dokumentu"}
            
            # Sprawdź, czy to jest PDF lub obraz
            if not doc.mime_type or (doc.mime_type != 'application/pdf' and not doc.mime_type.startswith('image/')):
                return {"error": "Ta funkcja obsługuje tylko pliki PDF i obrazy"}
            
            # Ścieżka do pliku
            file_path = FILES_DIR / doc.stored_filename
            if not file_path.exists():
                return {"error": "Nie znaleziono pliku"}
            
            # Obsługa PDF
            if doc.mime_type == 'application/pdf':
                # Pobierz liczbę stron z PDF
                try:
                    with open(file_path, 'rb') as pdf_file:
                        pdf_reader = PyPDF2.PdfReader(pdf_file)
                        total_pages = len(pdf_reader.pages)
                        
                        # Sprawdź, czy żądana strona istnieje
                        if page <= 0 or page > total_pages:
                            return {"error": f"Strona {page} nie istnieje. Dokument ma {total_pages} stron."}
                except Exception as e:
                    logger.error(f"Błąd odczytu dokumentu PDF: {str(e)}", exc_info=True)
                    return {"error": f"Nie można odczytać dokumentu PDF: {str(e)}"}
                
                # Sprawdź czy to jest zaznaczenie całej strony
                is_full_page = (abs(x1) < 0.01 and abs(y1) < 0.01 and abs(x2 - 1.0) < 0.01 and abs(y2 - 1.0) < 0.01)
                
                # Jeśli to zaznaczenie całej strony, sprawdź czy mamy już OCR
                if is_full_page:
                    ocr_txt_query = select(Document).where(
                        Document.ocr_parent_id == doc_id,
                        Document.doc_type == "OCR TXT"
                    )
                    ocr_txt = session.exec(ocr_txt_query).first()
                    
                    if ocr_txt:
                        # Mamy już OCR, zwróć go
                        from app.text_extraction import get_ocr_text_for_document
                        page_text = get_ocr_text_for_document(doc_id, session)
                        if page_text:
                            return {
                                "success": True,
                                "text": page_text.strip(),
                                "page": page,
                                "total_pages": total_pages,
                                "is_full_page": True
                            }
                
                # Konwertuj stronę PDF na obraz
                try:
                    # Konwertuj tylko wybraną stronę
                    images = convert_from_path(str(file_path), first_page=page, last_page=page, dpi=300)
                    
                    if not images:
                        return {"error": "Nie można skonwertować strony PDF na obraz"}
                    
                    # Weź pierwszy obraz strony
                    image = images[0]
                    
                except Exception as e:
                    logger.error(f"Błąd konwersji PDF na obraz: {str(e)}", exc_info=True)
                    return {"error": f"Błąd podczas konwersji PDF na obraz: {str(e)}"}
                    
            # Obsługa obrazów
            elif doc.mime_type.startswith('image/'):
                # Dla obrazów nie ma stron, zawsze używamy strony 1
                total_pages = 1
                
                # Sprawdź czy to jest zaznaczenie całego obrazu
                is_full_image = (abs(x1) < 0.01 and abs(y1) < 0.01 and abs(x2 - 1.0) < 0.01 and abs(y2 - 1.0) < 0.01)
                
                # Jeśli to zaznaczenie całego obrazu, sprawdź czy mamy już OCR
                if is_full_image:
                    ocr_txt_query = select(Document).where(
                        Document.ocr_parent_id == doc_id,
                        Document.doc_type == "OCR TXT"
                    )
                    ocr_txt = session.exec(ocr_txt_query).first()
                    
                    if ocr_txt:
                        # Mamy już OCR, zwróć go
                        from app.text_extraction import get_ocr_text_for_document
                        image_text = get_ocr_text_for_document(doc_id, session)
                        if image_text:
                            return {
                                "success": True,
                                "text": image_text.strip(),
                                "page": 1,
                                "total_pages": 1,
                                "is_full_image": True
                            }
                
                # Otwórz obraz bezpośrednio
                try:
                    image = Image.open(file_path)
                except Exception as e:
                    logger.error(f"Błąd otwarcia obrazu: {str(e)}", exc_info=True)
                    return {"error": f"Nie można otworzyć obrazu: {str(e)}"}
                    
            # Oblicz współrzędne zaznaczenia w pikselach
            width, height = image.size
            crop_x1 = int(x1 * width)
            crop_y1 = int(y1 * height)
            crop_x2 = int(x2 * width)
            crop_y2 = int(y2 * height)
            
            # Dodaj margines do zaznaczenia
            margin = 5
            crop_x1 = max(0, crop_x1 - margin)
            crop_y1 = max(0, crop_y1 - margin)
            crop_x2 = min(width, crop_x2 + margin)
            crop_y2 = min(height, crop_y2 + margin)
            
            # Wytnij zaznaczony fragment
            crop_image = image.crop((crop_x1, crop_y1, crop_x2, crop_y2))
            
            # Zapisz wycięty fragment do pliku tymczasowego
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp_path = tmp.name
            
            # Opcjonalne powiększenie małych fragmentów
            crop_width, crop_height = crop_image.size
            min_dimension = 300
            if crop_width < min_dimension or crop_height < min_dimension:
                scale_factor = min_dimension / min(crop_width, crop_height)
                new_width = int(crop_width * scale_factor)
                new_height = int(crop_height * scale_factor)
                crop_image = crop_image.resize((new_width, new_height), Image.LANCZOS)
            
            # Zapisz obraz fragmentu
            crop_image.save(tmp_path, format="PNG", quality=95)
            
            try:
                # Uruchom OCR na wyciętym fragmencie
                instruction = "Extract all the text visible in this image fragment. Keep all formatting."
                fragment_text = process_image_to_text(tmp_path, instruction=instruction)
                
                # Usuń plik tymczasowy
                import os
                os.unlink(tmp_path)
                
                # Zwróć wynik
                return {
                    "success": True,
                    "text": fragment_text.strip(),
                    "page": page,
                    "total_pages": total_pages
                }
                
            except Exception as e:
                logger.error(f"Błąd OCR fragmentu: {str(e)}", exc_info=True)
                return {
                    "success": True,
                    "text": "Nie udało się rozpoznać tekstu z fragmentu. Spróbuj zaznaczyć większy obszar.",
                    "page": page,
                    "total_pages": total_pages,
                    "error_fragment_ocr": str(e)
                }
    
    except Exception as e:
        logger.error(f"Globalny błąd OCR zaznaczenia: {str(e)}", exc_info=True)
        return {"error": f"Błąd: {str(e)}"}

@app.get("/document/{doc_id}/image-viewer")
def document_image_viewer(request: Request, doc_id: int):
    """Zaawansowany podgląd obrazu z funkcją zaznaczania i OCR."""
    with Session(engine) as session:
        doc = session.get(Document, doc_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Nie znaleziono dokumentu")
        
        # Sprawdź czy dokument to obraz
        if not doc.mime_type or not doc.mime_type.startswith('image/'):
            raise HTTPException(status_code=400, detail="Ten widok jest dostępny tylko dla plików obrazowych")
    
    return templates.TemplateResponse(
        "image_view_with_selection.html", 
        {
            "request": request, 
            "doc": doc,
            "title": f"Podgląd obrazu z zaznaczaniem - {doc.original_filename}"
        }
    )

# ==================== FUNKCJE POMOCNICZE ====================

async def _enqueue_ocr_documents_nonblocking(doc_ids: list[int]):
    """Asynchronicznie wstawia dokumenty do kolejki OCR bez blokowania."""
    from app.background_tasks import enqueue_ocr_task
    
    for doc_id in doc_ids:
        try:
            with Session(engine) as session:
                doc = session.get(Document, doc_id)
                if doc and doc.ocr_status == "pending":
                    await enqueue_ocr_task(doc_id)
                    await asyncio.sleep(0)
        except Exception as e:
            print(f"Błąd podczas dodawania dokumentu {doc_id} do kolejki OCR: {str(e)}")
            continue

def _generate_preview_html(request: Request, doc: Document, ocr_txt: Document = None) -> str:
    """Generuje HTML podglądu dokumentu - można wynieść do osobnego modułu."""
    # Przygotuj ścieżkę do pliku
    file_path = FILES_DIR / doc.stored_filename
    
    # Jeśli plik nie istnieje, zwróć błąd
    if not file_path.exists():
        return """
        <div class="alert alert-danger">
            <i class="bi bi-exclamation-triangle-fill me-2"></i>
            Plik nie istnieje.
        </div>
        """
    
    # Sprawdź typ pliku i przygotuj odpowiedni podgląd
    mime_type = doc.mime_type
    if not mime_type:
        mime_type = detect_mime_type(file_path)
    
    # Generuj odpowiednie URL-e
    download_url = str(request.url_for("document_download", doc_id=doc.id))
    
    # Dla Word (DOCX/DOC) wygeneruj podgląd tekstowy
    if mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                     'application/msword']:
        try:
            # Spróbuj użyć python-docx do odczytania zawartości
            from docx import Document as DocxDocument
            word_doc = DocxDocument(file_path)
            
            # Ekstrakcja tekstu
            text_content = []
            for para in word_doc.paragraphs:
                if para.text.strip():  # Pomijaj puste linie
                    text_content.append(para.text)
            
            # Jeśli nie ma treści, zwróć komunikat
            if not text_content:
                return f"""
                <div class="text-center py-4">
                    <i class="bi bi-file-earmark-word text-primary" style="font-size: 5rem;"></i>
                    <h5 class="mt-3">{doc.original_filename}</h5>
                    <p class="text-muted">Dokument Word jest pusty lub nie zawiera tekstu.</p>
                    <div class="mt-3">
                        <a href="{download_url}" class="btn btn-primary">
                            <i class="bi bi-download"></i> Pobierz dokument
                        </a>
                    </div>
                </div>
                """
            
            # Wygeneruj HTML z treścią dokumentu
            html_content = '<div class="p-4">'
            for paragraph in text_content:
                html_content += f'<p>{paragraph}</p>'
            html_content += '</div>'
            
            return f"""
            <div class="container-fluid">
                <div class="row mb-3">
                    <div class="col-12 text-center">
                        <i class="bi bi-file-earmark-word text-primary" style="font-size: 3rem;"></i>
                        <h4 class="mt-2">{doc.original_filename}</h4>
                        <p class="text-muted small">
                            Dokument Word • Dodano: {doc.upload_time.strftime('%Y-%m-%d %H:%M')}
                            {f" • Zmodyfikowano: {doc.last_modified.strftime('%Y-%m-%d %H:%M')}" if doc.last_modified else ""}
                        </p>
                    </div>
                </div>
                <div class="row">
                    <div class="col-12">
                        <div class="border rounded bg-white" style="max-height: 500px; overflow-y: auto;">
                            {html_content}
                        </div>
                    </div>
                </div>
                <div class="row mt-3">
                    <div class="col-12 text-center">
                        <a href="{download_url}" class="btn btn-primary">
                            <i class="bi bi-download"></i> Pobierz pełny dokument
                        </a>
                    </div>
                </div>
            </div>
            """
        except Exception as e:
            # W przypadku błędu, zwróć komunikat z informacją
            return f"""
            <div class="text-center py-4">
                <i class="bi bi-file-earmark-word text-primary" style="font-size: 5rem;"></i>
                <h5 class="mt-3">{doc.original_filename}</h5>
                <div class="alert alert-warning">
                    <i class="bi bi-exclamation-triangle-fill me-2"></i>
                    Nie można wygenerować podglądu dokumentu Word. Błąd: {str(e)}
                </div>
                <div class="mt-3">
                    <a href="{download_url}" class="btn btn-primary">
                        <i class="bi bi-download"></i> Pobierz dokument
                    </a>
                </div>
            </div>
            """

    # Dla PDF zwróć prosty widok
    elif mime_type == 'application/pdf':
        # Istniejący kod dla PDF
        pdf_url = str(request.url_for("document_preview", doc_id=doc.id))
        return f"""
        <div style="text-align: center; padding: 30px;">
            <div style="margin-bottom: 20px;">
                <i class="bi bi-file-earmark-pdf text-danger" style="font-size: 64px;"></i>
                <h4 style="margin-top: 15px;">{doc.original_filename}</h4>
                <p class="text-muted">
                    Dokument PDF • Dodano: {doc.upload_time.strftime('%Y-%m-%d %H:%M')}
                </p>
            </div>
            
            <div class="alert alert-info">
                <i class="bi bi-info-circle-fill me-2"></i>
                Podgląd PDF jest dostępny po otwarciu w nowej karcie.
            </div>
            
            <div style="margin-top: 20px;">
                <a href="{pdf_url}" class="btn btn-primary" target="_blank">
                    <i class="bi bi-eye"></i> Otwórz PDF
                </a>
                <a href="{download_url}" class="btn btn-outline-secondary ms-3">
                    <i class="bi bi-download"></i> Pobierz PDF
                </a>
            </div>
        </div>
        """

    # Dla innych typów zwróć podstawowy widok
    return f"""
    <div class="text-center py-4">
        <i class="bi bi-file-earmark" style="font-size: 5rem; color: #6c757d;"></i>
        <h5 class="mt-3">{doc.original_filename}</h5>
        <p class="text-muted">
            Typ pliku: {mime_type or "Nieznany"}<br>
            Dodano: {doc.upload_time.strftime('%Y-%m-%d %H:%M')}
        </p>
        <div class="mt-3">
            <a href="{download_url}" class="btn btn-primary">
                <i class="bi bi-download"></i> Pobierz dokument
            </a>
        </div>
    </div>
    """

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=80)
